from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime


def add_title_page(document: Document) -> None:
    title = document.add_paragraph()
    run = title.add_run("Functional Requirements Document (FRD)\nSchool Health EMR — Healthcare Management Module")
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("Version: 1.0").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Date: {datetime.utcnow().strftime('%Y-%m-%d')} (UTC)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Owner: School Health Program").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Prepared by: Project Team").alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = h
    for row in rows:
        row_cells = table.add_row().cells
        for idx, val in enumerate(row):
            row_cells[idx].text = val


def add_section_intro(document: Document) -> None:
    add_heading(document, "1. Purpose and Scope", 1)
    document.add_paragraph(
        "Deliver a user-friendly EMR module for school clinics to manage Individualized Health Care Plans (IHCP), document clinical encounters, administer medications, and monitor system-wide trends while complying with school health regulations and international standards.")

    add_heading(document, "2. In-Scope Modules", 1)
    add_bullets(document, [
        "Card 1: Individualized Health Care Plan (IHCP) Module",
        "Card 2: School Health Encounter Documentation",
        "Card 3: School Health Encounter Documentation Dashboard",
        "Card 4: Medication Administration Dashboard",
    ])

    add_heading(document, "3. Out of Scope (for this release)", 1)
    add_bullets(document, [
        "Parent portal submission flows (covered in Parent Forms module)",
        "External EHR/HIE integrations beyond secure document upload",
        "Telehealth/virtual visit features",
        "Inventory procurement and billing",
    ])

    add_heading(document, "4. Stakeholders and Roles", 1)
    add_table(document, ["Role", "Description", "Key Permissions"], [
        ["School Nurse", "Primary clinic operator", "Create/Update encounters, create IHCP, administer meds, view dashboards"],
        ["School Doctor", "Clinical oversight", "Approve IHCP, review encounters, advanced reporting"],
        ["School Admin", "Operational support", "View dashboards, manage rosters, non-clinical access"],
        ["Parent/Guardian", "Provide health data & consent via Parent Forms", "Submit/Update student medical info and consents"],
        ["DHA/Authority User", "Oversight & audit", "System-wide reporting, policy enforcement, audits"],
        ["System", "Automation engine", "Generate IDs, reminders, validations, alerts"],
    ])


def add_dependencies(document: Document) -> None:
    add_heading(document, "5. Dependencies on Other Solutions", 1)
    document.add_paragraph(
        "This module depends on the 'School Health – Student Detailed Medical Information (Parent Forms)' solution for authoritative student health data and consents.")

    add_heading(document, "5.1 Data Dependencies", 2)
    add_bullets(document, [
        "Student master data: Full Name, Student ID, Date of Birth, Gender, Grade/Class, Homeroom",
        "Medical profile: Diagnoses, Allergies, Chronic conditions, Disabilities, Medications, Family history, Special precautions",
        "Parent/Guardian details: Contacts and relationships",
        "Consents: Treatment, Medication administration, Emergency transfer, Data sharing",
        "Attachments: Specialist reports, action plans (e.g., asthma, anaphylaxis), lab results",
    ])

    add_heading(document, "5.2 Process Dependencies", 2)
    add_bullets(document, [
        "Annual Health File Review Notification: Triggers for current academic year to prompt verification",
        "Verification Workflow: Nurse/Doctor review and approval of parent-submitted updates",
        "Update Requests: Initiated by school health team to parents for missing or outdated info",
        "Data Cleansing: Periodic deduplication and normalization to maintain data quality",
    ])

    add_heading(document, "5.3 Technical Integration", 2)
    add_bullets(document, [
        "Shared Student ID and Academic Year dimensions",
        "Read-only access to verified parent-submitted data via secured APIs or shared data tables",
        "Event/Queue integration for status changes (e.g., verified, pending, rejected)",
        "Audit trail interoperability (user, timestamp, before/after values)",
    ])

    add_heading(document, "5.4 Dependency Risks & Mitigations", 2)
    add_table(document, ["Risk", "Impact", "Mitigation"], [
        ["Parent data not verified in time", "Delays IHCP creation and safe medication administration", "Escalation reminders; temporary nurse-entered notes flagged as unverified"],
        ["Consent not recorded or expired", "Cannot administer medication or perform procedures", "Block actions until consent; urgent override with incident log and parent notification"],
        ["Mismatch in Student IDs across systems", "Data linkage errors", "Master data sync jobs; reconciliation reports"],
        ["Outdated attachments", "Inaccurate care plans", "Annual reassessment alerts; required-document checks"],
    ])


def add_card1_ihcp(document: Document) -> None:
    add_heading(document, "6. Card 1 — Individualized Health Care Plan (IHCP)", 1)
    document.add_paragraph(
        "Objective: Create and maintain structured IHCPs for students with chronic conditions, linking diagnoses to ICD-10, defining goals, interventions, medications, accommodations, emergency protocols, and review schedules.")

    add_heading(document, "6.1 Core Features", 2)
    add_bullets(document, [
        "Create IHCP from verified student profile with auto-populated demographics",
        "Diagnosis mapping with ICD-10 and severity classification",
        "Care goals, interventions, expected outcomes with review dates",
        "Parent/Guardian consent linkage and validity checks",
        "Medication plan with dosing, route, times, required supplies, and school storage info",
        "Emergency action protocols (e.g., anaphylaxis, seizures) with quick-reference cards",
        "Accommodations/restrictions (e.g., PE restrictions, dietary needs)",
        "Automated reassessment alerts at start of academic year or upon new medical reports",
        "Document management: Upload and version specialist letters and action plans",
        "Audit trail for all edits and approvals",
    ])

    add_heading(document, "6.2 Data Model (Key Fields)", 2)
    add_table(document, ["Entity", "Field", "Type / Example", "Notes"], [
        ["IHCP", "IHCP ID", "Auto (IHCP-YYYY-####)", "Unique plan id"],
        ["IHCP", "Student ID", "Lookup", "From Student master"],
        ["IHCP", "Academic Year", "Text / 2025-2026", "Default current"],
        ["IHCP", "Diagnoses", "Multi-lookup ICD-10", "With severity"],
        ["IHCP", "Goals & Interventions", "Rich text", "Structured lists"],
        ["IHCP", "Medications", "Collection", "Name, dose, route, schedule, consent link"],
        ["IHCP", "Emergency Protocols", "Attachment/Template", "Condition-specific"],
        ["IHCP", "Accommodations", "Text", "PE/dietary/other"],
        ["IHCP", "Parent Consent", "Lookup", "Validate active/valid"],
        ["IHCP", "Status", "Draft/Active/Expired/Archived", "Workflow controlled"],
        ["IHCP", "Review Date", "Date", "Auto reminders"],
        ["IHCP", "Approver", "User", "Nurse/Doctor"],
    ])

    add_heading(document, "6.3 Workflow & Validations", 2)
    add_numbered(document, [
        "Draft -> Review -> Active -> Expired/Archived lifecycle",
        "Block activation if mandatory consents missing or expired",
        "Alert if medications lack administration instructions or storage details",
        "Auto-create tasks for annual reassessment and when parent uploads new reports",
        "Notify parents and homeroom when IHCP becomes Active (role-based)",
    ])

    add_heading(document, "6.4 Security & Access", 2)
    add_bullets(document, [
        "Only assigned nurse/doctor can edit; others read-only within same school",
        "Parents cannot edit IHCP; they may view published summary via portal if enabled",
        "All changes are audit-logged with user and timestamp",
    ])

    add_heading(document, "6.5 Reports & KPIs", 2)
    add_bullets(document, [
        "Active IHCPs by school/grade/condition",
        "IHCPs due for review in next 30/60/90 days",
        "Plans missing required consents or documents",
        "Outcome tracking against goals (qualitative/quantitative)",
    ])


def add_card2_encounters(document: Document) -> None:
    add_heading(document, "7. Card 2 — School Health Encounter Documentation", 1)
    document.add_paragraph(
        "Objective: Record and monitor all clinic visits—routine complaints, injuries, emergencies, and medication administration—with standardized data capture, classifications, and outcomes.")

    add_heading(document, "7.1 Core Features", 2)
    add_bullets(document, [
        "Auto-capture demographics and visit details (Visit Number, Date/Time, Academic Year, Staff)",
        "Vital signs and assessments with reference ranges by age",
        "Structured presenting complaints via dropdowns with free-text notes",
        "Incident and medication error documentation",
        "Immediate and follow-up actions logging with tasks and reminders",
        "Visit classification (illness, injury, first aid, emergency, medication admin)",
        "Standardized outcomes (returned to class, sent home, ER transfer, follow-up)",
        "Attachments: photos of injuries, referral letters",
    ])

    add_heading(document, "7.2 Data Model (Key Fields)", 2)
    add_table(document, ["Entity", "Field", "Type / Example", "Notes"], [
        ["Encounter", "Visit Number", "Auto (ENC-YYYY-####)", "Unique encounter id"],
        ["Encounter", "Student ID", "Lookup", "From Student master"],
        ["Encounter", "Date/Time", "Datetime", "Auto-filled"],
        ["Encounter", "Staff", "User", "Auto-filled"],
        ["Encounter", "Chief Complaint", "Picklist + text", "Configurable"],
        ["Encounter", "Vital Signs", "Nested fields", "BP, HR, RR, Temp, SpO2, Pain"],
        ["Encounter", "Assessment", "Text", "Clinical notes"],
        ["Encounter", "Actions/Procedures", "Multi-select", "First aid, meds, referral"],
        ["Encounter", "Outcome", "Picklist", "Returned, Sent home, ER, Other"],
        ["Encounter", "Follow-up", "Task/Date", "Optional"],
        ["Encounter", "Incident/Medication Error", "Boolean + details", "If applicable"],
    ])

    add_heading(document, "7.3 Workflow & Validations", 2)
    add_numbered(document, [
        "Create -> Save -> Complete lifecycle",
        "Require outcome before completion",
        "Block medication administration without active consent unless emergency override with incident log",
        "Auto-notify parent for significant events (injury, ER transfer)",
        "Generate follow-up tasks when indicated",
    ])

    add_heading(document, "7.4 Security & Access", 2)
    add_bullets(document, [
        "Editable by encounter creator and assigned clinic staff; read-only to same-school clinicians",
        "Parents receive notifications/summary per policy but cannot view full clinical notes",
        "Full audit history of changes",
    ])

    add_heading(document, "7.5 Reports & KPIs", 2)
    add_bullets(document, [
        "Visits per day/week/month by classification",
        "Injury patterns by location/activity",
        "Incidents/medication errors with root-cause tagging",
        "Time-to-disposition and follow-up completion rates",
    ])


def add_card3_dashboard(document: Document) -> None:
    add_heading(document, "8. Card 3 — School Health Encounter Documentation Dashboard", 1)
    document.add_paragraph(
        "Objective: Provide real-time operational visibility into encounters, triage, and outcomes across schools, highlighting trends and overdue items.")

    add_heading(document, "8.1 Widgets & Visualizations", 2)
    add_bullets(document, [
        "Live encounters today by status",
        "Visits by type (illness, injury, first aid, emergency)",
        "Top presenting complaints",
        "Outcomes distribution",
        "Follow-ups due/overdue",
        "Heatmap by grade/class",
        "Incidents and medication errors trend",
    ])

    add_heading(document, "8.2 Filters & Drilldowns", 2)
    add_bullets(document, [
        "Date range, school, grade, staff, classification",
        "Click-through to encounter records",
        "Export to CSV/PDF",
    ])

    add_heading(document, "8.3 Security", 2)
    add_bullets(document, [
        "Row-level security by school",
        "De-identified views for broader stakeholders if required",
    ])


def add_card4_med_dashboard(document: Document) -> None:
    add_heading(document, "9. Card 4 — Medication Administration Dashboard", 1)
    document.add_paragraph(
        "Objective: Monitor scheduled and PRN medication administrations, consents, stock status, missed doses, and errors to enhance safety and compliance.")

    add_heading(document, "9.1 Core Features", 2)
    add_bullets(document, [
        "Daily MAR (Medication Administration Record) view with due/administered/missed",
        "Consent validation flags and expiry alerts",
        "Integration with IHCP medication plans",
        "Exception logging: missed doses, refusals, errors",
        "Inventory snapshot (on-hand, expiring soon)",
        "Notifications to parents for administered/ missed critical doses per policy",
    ])

    add_heading(document, "9.2 Measures & KPIs", 2)
    add_bullets(document, [
        "Administration adherence rate",
        "Missed/late doses per 1,000 administrations",
        "Error rate and categories",
        "Stockouts and near-expiry items",
    ])


def add_non_functional(document: Document) -> None:
    add_heading(document, "10. Non-Functional Requirements", 1)
    add_bullets(document, [
        "Availability: 99.5% during school hours",
        "Performance: <2s for record open/save under normal load",
        "Security: Role-based access, field-level security for sensitive data",
        "Privacy: Compliance with applicable student data protection regulations",
        "Auditability: Immutable audit logs for all clinical changes",
        "Interoperability: Standards-ready (ICD-10, CSV export)",
        "Accessibility: WCAG 2.1 AA for dashboards and forms",
        "Localization: Multi-language labels (EN/AR) where applicable",
    ])


def add_acceptance_criteria(document: Document) -> None:
    add_heading(document, "11. Acceptance Criteria (High-Level)", 1)
    add_numbered(document, [
        "IHCP can be created, approved, and activated only when required consents are valid",
        "Encounter completion requires outcome and saves audit trail",
        "Medication admin cannot be logged without consent unless emergency override with incident",
        "Dashboards reflect new data within 5 minutes",
        "Annual reassessment tasks auto-generate for active IHCPs at academic year start",
    ])


def add_traceability(document: Document) -> None:
    add_heading(document, "12. Requirements Traceability (BRD → FRD)", 1)
    add_table(document, ["BRD Statement", "FRD Section"], [
        ["Create IHCP with ICD-10 and parent consents", "6. Card 1 — IHCP: 6.1, 6.2, 6.3"],
        ["Document and classify all clinic visits", "7. Card 2 — Encounters: 7.1, 7.2"],
        ["Provide encounter dashboards", "8. Card 3 — Dashboard: 8.1, 8.2"],
        ["Monitor medication administration and errors", "9. Card 4 — Medication Dashboard: 9.1, 9.2"],
        ["Integrate with Parent Forms module", "5. Dependencies: 5.1–5.4"],
    ])


def add_security_compliance(document: Document) -> None:
    add_heading(document, "13. Security, Compliance, and Audit", 1)
    add_bullets(document, [
        "Role-based access control with least privilege",
        "Field-level protection for sensitive diagnoses and notes",
        "Comprehensive audit trails (create, update, delete, view) including user and timestamp",
        "Data retention and archival aligned with school authority policy",
        "Incident reporting workflow for privacy and safety events",
    ])


def add_appendices(document: Document) -> None:
    add_heading(document, "Appendix A — Glossary", 1)
    add_table(document, ["Term", "Definition"], [
        ["IHCP", "Individualized Health Care Plan"],
        ["ICD-10", "International Classification of Diseases, 10th Revision"],
        ["MAR", "Medication Administration Record"],
    ])

    add_heading(document, "Appendix B — Sample Field-Level Specs (Abbreviated)", 1)
    add_table(document, ["Entity", "Field", "Type", "Validation"], [
        ["Encounter", "Outcome", "Picklist", "Required on Complete"],
        ["IHCP", "Parent Consent", "Lookup", "Must be Active to activate IHCP"],
        ["Medication Admin", "Consent Check", "Boolean/Auto", "Block without consent unless emergency override"],
    ])


if __name__ == "__main__":
    doc = Document()
    core_props = doc.core_properties
    core_props.title = "FRD — School Health EMR: Healthcare Management Module"
    core_props.author = "School Health Program"
    core_props.subject = "Functional Requirements Document"
    core_props.created = datetime.utcnow()

    add_title_page(doc)
    add_section_intro(doc)
    add_dependencies(doc)
    add_card1_ihcp(doc)
    add_card2_encounters(doc)
    add_card3_dashboard(doc)
    add_card4_med_dashboard(doc)
    add_non_functional(doc)
    add_acceptance_criteria(doc)
    add_traceability(doc)
    add_security_compliance(doc)
    add_appendices(doc)

    output_path = "/workspace/FRD_School_Health_EMR_Healthcare_Management_Module.docx"
    doc.save(output_path)
    print(f"Generated: {output_path}")